from typing import Optional  # Importiert das Modul für optionale Typen

import os  # Importiert das Modul für Betriebssystem-interaktive Funktionen
from docx import Document  # Importiert das Modul zur Arbeit mit Word-Dokumenten
from docx2pdf import convert  
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx.shared import Inches, Pt  
from matplotlib.figure import Figure  # Importiert die Figur-Klasse für die Darstellung von Diagrammen

# Artem
def save_as_pdf(doc_path: str, output_dir: str):  # Funktion zur Konvertierung von .docx nach .pdf
    # Konvertiert das .docx-Dokument in ein PDF
    pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".pdf")  # Setzt den Pfad für das PDF
    convert(doc_path, output_path=output_dir)  # Führt die Konvertierung durch
    print(f"PDF erstellt: {pdf_path}")  # Gibt den Speicherort des erstellten PDFs aus

    return pdf_path  # Gibt den Pfad der PDF-Datei zurück

def create_report_template():  # Funktion zur Erstellung einer Berichtsvorlage
    # Stellt sicher, dass das Verzeichnis für Vorlagen existiert
    templates_dir = os.path.join(".", "templates")
    if not os.path.exists(templates_dir):  # Wenn das Verzeichnis nicht existiert
        os.makedirs(templates_dir)  # Erstellt das Verzeichnis

    # Erstellt ein neues Dokument
    doc = Document()

    # Fügt den Titel hinzu
    title = doc.add_heading(level=1)  # Fügt eine Überschrift 1 hinzu
    title_run = title.add_run("Bericht [week]")  # Fügt Text in die Überschrift ein
    title_run.font.size = Pt(16)  # Setzt die Schriftgröße auf 16 Punkte
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Zentriert den Text

    # Fügt Abschnitte mit Platzhaltertext hinzu
    sections = [
        "Wochenbericht",  # Abschnittsname
        "Vergleich zur vorherigen Woche",  # Abschnittsname
    ]

    for section in sections:  # Iteriert über die Abschnitte
        doc.add_heading(section, level=2)  # Fügt jede Abschnittsüberschrift hinzu
        doc.add_paragraph(f"[{section.lower().replace(' ', '')}]")  # Fügt Platzhaltertext hinzu

    # Fügt einen Platzhalter für mögliche Bilder hinzu
    doc.add_paragraph("[image]").alignment = WD_ALIGN_PARAGRAPH.CENTER  # Fügt Platzhalter für Bild hinzu und zentriert es

    # Baut den Pfad für die Vorlage zusammen
    template_path = os.path.join(templates_dir, "report_template.docx")
    doc.save(template_path)  # Speichert die Vorlage

    print(f"Vorlage gespeichert unter: {os.path.abspath(template_path)}")  # Gibt den absoluten Pfad der Vorlage aus


def insert_content(  # Funktion zum Einfügen von Inhalt in die Vorlage
    content_dict: dict[str, str],  # Dictionary mit den Inhalten
    image: Figure,  # Matplotlib-Figur für das Bild
    template_path: Optional[str] = r".\templates\report_template.docx",  # Pfad zur Vorlage
) -> str:
    """
    Fügt tatsächlichen Inhalt in die Vorlage ein

    content_dict sollte ein Dictionary mit Schlüsseln enthalten, die den Abschnittsnamen entsprechen,
    und Werten, die den Text und die Bildpfade enthalten
    """
    doc = Document(template_path)  # Lädt das Dokument basierend auf der Vorlage

    # Ersetzt den Platzhalter für die Woche im Titel
    for paragraph in doc.paragraphs:  # Iteriert über alle Absätze im Dokument
        if "[week]" in paragraph.text:  # Wenn der Platzhalter für die Woche gefunden wird
            for run in paragraph.runs:  # Iteriert über alle Textteile im Absatz
                run.text = run.text.replace("[week]", content_dict.get("week", ""))  # Ersetzt den Platzhalter mit dem Wert

    # Ersetzt Platzhalterinhalte
    for paragraph in doc.paragraphs:  # Iteriert über alle Absätze im Dokument
        if "[wochenbericht]" in paragraph.text:  # Wenn der Platzhalter für Wochenbericht gefunden wird
            paragraph.text = content_dict.get("Wochenbericht", "")  # Ersetzt mit dem Inhalt des Wochenberichts
        elif "[vergleichzurvorherigenwoche]" in paragraph.text:  # Wenn der Platzhalter für Vergleich gefunden wird
            paragraph.text = content_dict.get("Vergleich zur vorherigen Woche", "")  # Ersetzt mit dem Vergleichstext

        # Ersetzt den Bildplatzhalter mit dem tatsächlichen Bild
        if "[image]" in paragraph.text:  # Wenn der Bildplatzhalter gefunden wird
            paragraph.text = ""  # Löscht den Platzhaltertext
            # Speichert das Matplotlib-Diagramm als temporäre Datei
            image.savefig("temp_plot.png")
            # Fügt das Bild zum Dokument hinzu
            doc.add_picture("temp_plot.png", width=Inches(6))  # Bildbreite wird auf 6 Inches gesetzt
            # Entfernt die temporäre Datei
            import os
            os.remove("temp_plot.png")

    # Speichert das fertige Dokument
    doc.save("report_final.docx")

    doc_path = "report_final.docx"  # Pfad zum gespeicherten Dokument
    output_dir = os.getcwd()  # Aktuelles Arbeitsverzeichnis
    pdf_path = save_as_pdf(doc_path, output_dir)  # Konvertiert das Dokument in eine PDF-Datei

    os.remove("report_final.docx")  # Löscht die temporäre .docx-Datei
    os.startfile(pdf_path)  # Öffnet die PDF-Datei
