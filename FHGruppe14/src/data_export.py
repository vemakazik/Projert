from typing import Optional

import os
from docx import Document
from docx2pdf import convert #Artem
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib.figure import Figure

#Artem
def save_as_pdf(doc_path: str, output_dir: str):
    # Converting .docx -> .pdf
    pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".pdf")
    convert(doc_path, output_path=output_dir)
    print(f"PDF created: {pdf_path}")

    return pdf_path # could be deleted

def create_report_template():
    # Ensure templates directory exists
    templates_dir = os.path.join(".", "templates")
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)

    # Create new document
    doc = Document()

    # Add title
    title = doc.add_heading(level=1)
    title_run = title.add_run("Bericht [week]")
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add sections with placeholder text
    sections = [
        "Wochenbericht",
        "Vergleich zur vorherigen Woche",
    ]

    for section in sections:
        doc.add_heading(section, level=2)
        doc.add_paragraph(f"[{section.lower().replace(' ', '')}]")

    # Add placeholder for potential images
    doc.add_paragraph("[image]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Use os.path.join to build the save path correctly
    template_path = os.path.join(templates_dir, "report_template.docx")
    doc.save(template_path)

    print(f"Template saved at: {os.path.abspath(template_path)}")


def insert_content(
    content_dict: dict[str, str],
    image: Figure,
    template_path: Optional[str] = r".\templates\report_template.docx",
) -> str:
    """
    Insert actual content into the template

    content_dict should be a dictionary with keys matching section names
    and values containing the text and image paths
    """
    doc = Document(template_path)

    # Replace the week placeholder in title
    for paragraph in doc.paragraphs:
        if "[week]" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("[week]", content_dict.get("week", ""))

    # Replace content placeholders
    for paragraph in doc.paragraphs:
        if "[wochenbericht]" in paragraph.text:
            paragraph.text = content_dict.get("Wochenbericht", "")
        elif "[vergleichzurvorherigenwoche]" in paragraph.text:
            paragraph.text = content_dict.get("Vergleich zur vorherigen Woche", "")

        # Replace image placeholder with actual image
        if "[image]" in paragraph.text:
            paragraph.text = ""  # Clear the placeholder text
            # Save the matplotlib figure to a temporary file
            image.savefig("temp_plot.png")
            # Add the image to the document
            doc.add_picture("temp_plot.png", width=Inches(6))  # Adjust width as needed
            # Remove temporary file
            import os

            os.remove("temp_plot.png")

   #Artem
    doc.save("report_final.docx")

    doc_path = "report_final.docx"
    output_dir = os.getcwd()
    pdf_path = save_as_pdf(doc_path, output_dir)

    os.remove("report_final.docx")
    os.startfile(pdf_path)